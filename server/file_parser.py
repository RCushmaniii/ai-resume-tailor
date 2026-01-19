"""
File Parser Module

Extracts text from PDF and DOCX resume files.
Processes files in memory without storing them.

File: server/file_parser.py
"""

import io
import logging
from typing import Optional, Tuple, List

# Configure logging
logger = logging.getLogger(__name__)

# Constants
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
MIN_TEXT_LENGTH = 100  # Minimum characters for a valid resume
MIN_CHARS_PER_PAGE = 50  # Minimum characters expected per page (to detect image-based pages)
SUPPORTED_EXTENSIONS = {'pdf', 'docx'}


class FileParserError(Exception):
    """Custom exception for file parsing errors."""
    pass


def extract_text_from_pdf(file_bytes: bytes) -> dict:
    """
    Extract text from PDF file bytes.

    Args:
        file_bytes: Raw PDF file content

    Returns:
        Dictionary with:
        - text: Extracted text from all pages
        - total_pages: Total number of pages in PDF
        - pages_with_text: Number of pages that had extractable text
        - warning: Optional warning message if extraction was partial

    Raises:
        FileParserError: If PDF parsing fails completely
    """
    try:
        import pdfplumber
    except ImportError:
        raise FileParserError("PDF parsing library not installed. Please install pdfplumber.")

    text_parts = []
    total_pages = 0
    pages_with_text = 0
    pages_with_minimal_text = []  # Pages that had very little text (likely image-based)

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            total_pages = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        stripped_text = page_text.strip()
                        if stripped_text:
                            text_parts.append(stripped_text)
                            if len(stripped_text) >= MIN_CHARS_PER_PAGE:
                                pages_with_text += 1
                            else:
                                pages_with_minimal_text.append(page_num)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
                    continue

    except Exception as e:
        logger.error(f"PDF parsing error: {e}")
        raise FileParserError(f"Failed to parse PDF: {str(e)}")

    combined_text = "\n\n".join(text_parts)

    # Determine if there's a warning to show
    warning = None
    if total_pages > 0 and pages_with_text == 0:
        # No pages had meaningful text - likely all image-based
        warning = "image_based"
    elif pages_with_minimal_text and len(pages_with_minimal_text) > total_pages / 2:
        # More than half the pages had minimal text
        warning = "partial_extraction"
    elif total_pages > pages_with_text + len(pages_with_minimal_text):
        # Some pages completely failed to extract
        warning = "some_pages_failed"

    return {
        "text": combined_text,
        "total_pages": total_pages,
        "pages_with_text": pages_with_text,
        "pages_with_minimal_text": len(pages_with_minimal_text),
        "warning": warning
    }


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX file bytes.

    Args:
        file_bytes: Raw DOCX file content

    Returns:
        Extracted text from document

    Raises:
        FileParserError: If DOCX parsing fails
    """
    try:
        from docx import Document
    except ImportError:
        raise FileParserError("DOCX parsing library not installed. Please install python-docx.")

    text_parts = []

    try:
        doc = Document(io.BytesIO(file_bytes))

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        # Extract text from tables (common in resumes)
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

    except Exception as e:
        logger.error(f"DOCX parsing error: {e}")
        raise FileParserError(f"Failed to parse DOCX: {str(e)}")

    return "\n".join(text_parts)


def validate_file(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    """
    Validate file size and type.

    Args:
        file_bytes: Raw file content
        filename: Original filename

    Returns:
        Tuple of (extension, mime_type)

    Raises:
        FileParserError: If validation fails
    """
    # Check file size
    if len(file_bytes) > MAX_FILE_SIZE:
        size_mb = MAX_FILE_SIZE // (1024 * 1024)
        raise FileParserError(f"File too large. Maximum size is {size_mb}MB.")

    if len(file_bytes) == 0:
        raise FileParserError("File is empty.")

    # Get and validate extension
    if '.' not in filename:
        raise FileParserError("File must have an extension (.pdf or .docx).")

    extension = filename.lower().rsplit('.', 1)[-1]

    if extension not in SUPPORTED_EXTENSIONS:
        raise FileParserError(
            f"Unsupported file format: .{extension}. "
            f"Please use PDF or DOCX files."
        )

    return extension


def parse_resume_file(file_bytes: bytes, filename: str) -> dict:
    """
    Parse resume file and extract text.

    This is the main entry point for file parsing.

    Args:
        file_bytes: Raw file bytes
        filename: Original filename (for extension detection)

    Returns:
        Dictionary with:
        - text: Extracted text content
        - character_count: Length of extracted text
        - file_type: Type of file processed
        - warning: Optional warning code for partial extraction
        - warning_message: Human-readable warning message (if applicable)

    Raises:
        FileParserError: If parsing fails completely
    """
    # Validate file
    extension = validate_file(file_bytes, filename)

    # Extract text based on file type
    warning = None
    warning_message = None

    if extension == 'pdf':
        pdf_result = extract_text_from_pdf(file_bytes)
        text = pdf_result["text"]
        warning = pdf_result.get("warning")
        file_type = 'application/pdf'

        # Handle image-based PDF (no text at all)
        if warning == "image_based" and not text.strip():
            raise FileParserError(
                "This PDF appears to be image-based (like a scanned document) "
                "and we couldn't extract the text. Please try one of these options:\n\n"
                "1. Copy and paste your resume text directly into the text field\n"
                "2. If you have the original Word document, upload that instead\n"
                "3. Use a PDF with selectable text"
            )

        # Generate warning messages for partial extraction
        if warning == "partial_extraction":
            warning_message = (
                f"We extracted text from your PDF, but some pages may contain "
                f"images or graphics that couldn't be read. Please review the "
                f"extracted text below to make sure it looks complete."
            )
        elif warning == "some_pages_failed":
            warning_message = (
                f"We extracted text from {pdf_result['pages_with_text']} of "
                f"{pdf_result['total_pages']} pages. Some pages may have failed "
                f"to process. Please review the extracted text to ensure nothing "
                f"important is missing."
            )

    elif extension == 'docx':
        text = extract_text_from_docx(file_bytes)
        file_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    else:
        raise FileParserError(f"Unsupported file type: {extension}")

    # Clean up text
    text = text.strip()

    # Validate extracted content
    if not text:
        raise FileParserError(
            "We couldn't extract any text from your file. This sometimes happens with:\n\n"
            "- Scanned documents or image-based PDFs\n"
            "- Files with unusual formatting\n"
            "- Corrupted or password-protected files\n\n"
            "Please try copying and pasting your resume text directly into the text field instead."
        )

    if len(text) < MIN_TEXT_LENGTH:
        raise FileParserError(
            f"The extracted text seems too short ({len(text)} characters). "
            f"A typical resume has at least {MIN_TEXT_LENGTH} characters.\n\n"
            "This might mean:\n"
            "- Only part of your resume was extracted\n"
            "- The file contains mostly images or graphics\n\n"
            "Please try copying and pasting your full resume text instead."
        )

    result = {
        'text': text,
        'character_count': len(text),
        'file_type': file_type
    }

    # Add warning info if applicable
    if warning and warning_message:
        result['warning'] = warning
        result['warning_message'] = warning_message

    return result


# Quick test function for development
if __name__ == '__main__':
    import sys

    if len(sys.argv) < 2:
        print("Usage: python file_parser.py <filename>")
        sys.exit(1)

    filename = sys.argv[1]

    try:
        with open(filename, 'rb') as f:
            file_bytes = f.read()

        result = parse_resume_file(file_bytes, filename)
        print(f"File type: {result['file_type']}")
        print(f"Character count: {result['character_count']}")
        print(f"\n--- Extracted Text ---\n")
        print(result['text'][:2000])
        if len(result['text']) > 2000:
            print(f"\n... (truncated, total {result['character_count']} chars)")

    except FileParserError as e:
        print(f"Error: {e}")
        sys.exit(1)
    except FileNotFoundError:
        print(f"File not found: {filename}")
        sys.exit(1)
